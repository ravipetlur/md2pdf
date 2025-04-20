from flask import Flask, render_template, request, send_file
import markdown
import tempfile
import os
from weasyprint import HTML, CSS
from docx import Document
from io import BytesIO
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re

app = Flask(__name__)

def convert_to_pdf(markdown_content):
    # Convert markdown to HTML
    html_content = markdown.markdown(markdown_content, extensions=['tables', 'fenced_code'])
    
    # Add comprehensive styling
    styled_html = f"""
    <html>
        <head>
            <style>
                @page {{
                    size: A4;
                    margin: 2.5cm 1.5cm;
                }}
                
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    font-size: 11pt;
                    max-width: 100%;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    margin: 0;
                    padding: 0;
                }}
                
                p {{
                    margin: 1em 0;
                    max-width: 100%;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }}
                
                h1, h2, h3, h4, h5, h6 {{
                    margin-top: 1.5em;
                    margin-bottom: 0.5em;
                    line-height: 1.2;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }}
                
                code {{
                    background-color: #f4f4f4;
                    padding: 2px 5px;
                    border-radius: 3px;
                    font-family: 'Courier New', monospace;
                    font-size: 0.9em;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    white-space: pre-wrap;
                }}
                
                pre {{
                    background-color: #f4f4f4;
                    padding: 15px;
                    border-radius: 5px;
                    overflow-x: auto;
                    white-space: pre-wrap;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    max-width: 100%;
                    margin: 1em 0;
                }}
                
                pre code {{
                    background-color: transparent;
                    padding: 0;
                    border-radius: 0;
                    white-space: pre-wrap;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }}
                
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 1em 0;
                    page-break-inside: auto;
                }}
                
                tr {{
                    page-break-inside: avoid;
                    page-break-after: auto;
                }}
                
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    max-width: 300px;
                }}
                
                th {{
                    background-color: #f4f4f4;
                }}
                
                img {{
                    max-width: 100%;
                    height: auto;
                }}
                
                ul, ol {{
                    padding-left: 2em;
                    margin: 1em 0;
                }}
                
                li {{
                    margin: 0.5em 0;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }}
                
                blockquote {{
                    margin: 1em 0;
                    padding-left: 1em;
                    border-left: 4px solid #ddd;
                    color: #666;
                }}
                
                /* Ensure page breaks are handled gracefully */
                h1, h2, h3, h4, h5, h6 {{
                    page-break-after: avoid;
                    page-break-inside: avoid;
                }}
                
                img, table, figure {{
                    page-break-inside: avoid;
                }}
                
                /* Handle long URLs and strings */
                a {{
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    white-space: pre-wrap;
                    color: #0066cc;
                    text-decoration: underline;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
    </html>
    """
    
    # Create PDF using WeasyPrint
    pdf_file = BytesIO()
    HTML(string=styled_html).write_pdf(
        target=pdf_file,
        stylesheets=[CSS(string='@page { size: A4; margin: 2.5cm 1.5cm; }')]
    )
    pdf_file.seek(0)
    return pdf_file

def convert_to_docx(markdown_content):
    # Convert markdown to HTML first
    html_content = markdown.markdown(markdown_content, extensions=['tables', 'fenced_code'])
    
    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create a new Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    def add_heading(text, level):
        heading = doc.add_heading(text, level=level)
        heading.style.font.name = 'Arial'
        return heading

    def add_code_block(text):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        return p

    def add_list_item(text, is_ordered, level=0):
        # Create paragraph with appropriate indentation
        p = doc.add_paragraph()
        # Set indentation based on level (0.5 inches per level)
        p.paragraph_format.left_indent = Inches(0.5 * level)
        
        # Add bullet/number
        if is_ordered:
            # For ordered lists, manually add the number
            p.add_run(f"{level+1}. ")
        else:
            # For unordered lists, use bullet point
            p.add_run("• ")
        
        # Add the actual text
        p.add_run(text)
        return p

    # Process each element in the HTML
    for element in soup.descendants:
        if isinstance(element, str) or element.name is None:
            continue
            
        if element.name == 'h1':
            add_heading(element.get_text().strip(), 0)
        elif element.name == 'h2':
            add_heading(element.get_text().strip(), 1)
        elif element.name == 'h3':
            add_heading(element.get_text().strip(), 2)
        elif element.name == 'h4':
            add_heading(element.get_text().strip(), 3)
        elif element.name == 'p':
            if element.find('code'):
                # Handle inline code
                p = doc.add_paragraph()
                for item in element.children:
                    if item.name == 'code':
                        run = p.add_run(item.get_text())
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    else:
                        p.add_run(str(item))
            else:
                text = element.get_text().strip()
                if text:  # Only add non-empty paragraphs
                    doc.add_paragraph(text)
        elif element.name == 'pre':
            # Handle code blocks
            code_text = element.get_text().strip()
            if code_text:
                add_code_block(code_text)
        elif element.name == 'ul':
            for i, li in enumerate(element.find_all('li', recursive=False)):
                # Calculate nesting level based on parent ul/ol elements
                level = len([p for p in li.parents if p.name in ['ul', 'ol']])
                add_list_item(li.get_text().strip(), False, level - 1)
        elif element.name == 'ol':
            for i, li in enumerate(element.find_all('li', recursive=False)):
                # Calculate nesting level based on parent ul/ol elements
                level = len([p for p in li.parents if p.name in ['ul', 'ol']])
                add_list_item(li.get_text().strip(), True, level - 1)
        elif element.name == 'table':
            # Handle tables
            table_rows = element.find_all('tr')
            if table_rows:
                num_cols = max(len(row.find_all(['td', 'th'])) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                
                for i, row in enumerate(table_rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text().strip()
                        
                        # Make header row bold
                        if i == 0 or cell.name == 'th':
                            for paragraph in table.cell(i, j).paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
            
            doc.add_paragraph()  # Add spacing after table

    # Save to BytesIO
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Check if file was uploaded
        if 'file' not in request.files:
            return 'No file uploaded', 400
        
        file = request.files['file']
        if file.filename == '':
            return 'No file selected', 400
        
        # Read the markdown content
        markdown_content = file.read().decode('utf-8')
        output_format = request.form.get('format', 'pdf')
        
        if output_format == 'pdf':
            pdf_file = convert_to_pdf(markdown_content)
            return send_file(
                pdf_file,
                mimetype='application/pdf',
                as_attachment=True,
                download_name='converted.pdf'
            )
        else:
            docx_file = convert_to_docx(markdown_content)
            return send_file(
                docx_file,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name='converted.docx'
            )
    
    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)
