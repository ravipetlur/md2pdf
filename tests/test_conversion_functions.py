import pytest
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from app import convert_to_pdf, convert_to_docx
from docx import Document
from PyPDF2 import PdfReader

@pytest.fixture
def sample_markdown():
    return """# Test Document

This is a test paragraph with **bold** and *italic* text.

## Code Example
```python
def test_function():
    return "Hello, World!"
```

## List Example
1. First item
2. Second item
   - Nested item
   - Another nested item

## Table Example
| Header 1 | Header 2 |
|----------|----------|
| Cell 1   | Cell 2   |
"""

def test_pdf_conversion_function(sample_markdown):
    """Test the PDF conversion function directly"""
    pdf_file = convert_to_pdf(sample_markdown)
    
    # Verify PDF was created
    assert pdf_file is not None
    
    # Read PDF content
    pdf_reader = PdfReader(pdf_file)
    text = pdf_reader.pages[0].extract_text()
    
    # Verify content
    assert 'Test Document' in text
    assert 'test paragraph' in text
    assert 'Code Example' in text
    assert 'test_function' in text
    assert 'First item' in text
    assert 'Header 1' in text

def test_docx_conversion_function(sample_markdown):
    """Test the DOCX conversion function directly"""
    docx_file = convert_to_docx(sample_markdown)
    
    # Verify DOCX was created
    assert docx_file is not None
    
    # Read DOCX content
    doc = Document(docx_file)
    full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    
    # Verify content
    assert 'Test Document' in full_text
    assert 'test paragraph' in full_text
    assert 'Code Example' in full_text
    assert 'test_function' in full_text
    assert 'First item' in full_text

def test_special_characters():
    """Test handling of special characters"""
    special_markdown = """# Special Characters Test

Text with special characters: áéíóú ñ & < > " '

```
Special chars in code: © ® ™ € £ ¥
```
"""
    # Test PDF
    pdf_file = convert_to_pdf(special_markdown)
    pdf_reader = PdfReader(pdf_file)
    text = pdf_reader.pages[0].extract_text()
    assert 'Special Characters Test' in text
    assert 'áéíóú' in text
    assert '© ® ™ € £ ¥' in text

    # Test DOCX
    docx_file = convert_to_docx(special_markdown)
    doc = Document(docx_file)
    full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    assert 'Special Characters Test' in full_text
    assert 'áéíóú' in full_text
    assert '© ® ™ € £ ¥' in full_text